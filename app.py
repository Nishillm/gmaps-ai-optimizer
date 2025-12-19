import streamlit as st
import pandas as pd
from google import genai
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium_stealth import stealth
from docx import Document
import io
import time

# --- UI THEME ---
st.set_page_config(page_title="Hunter Pro Dashboard", layout="wide")
st.markdown("""
    <style>
    .stApp { background-color: #FFFFFF; }
    [data-testid="stSidebar"] { background-color: #FCE4EC !important; }
    div[data-baseweb="input"] > div, div[data-baseweb="textarea"] > div {
        background-color: #E1BEE7 !important; border-radius: 8px !important;
    }
    input, textarea { color: #000000 !important; font-weight: 500; }
    label, p, h1, h2, h3, .stMarkdown, span { color: #000000 !important; }
    .stButton>button {
        background-color: #9575CD; color: white; border-radius: 8px; width: auto; font-weight: bold;
    }
    </style>
    """, unsafe_allow_html=True)

# --- SESSION STATE ---
if 'history' not in st.session_state: st.session_state['history'] = []
if 'leads' not in st.session_state: st.session_state['leads'] = []

# --- API CLIENT SETUP ---
try:
    # Ensure GEMINI_API_KEY is in your Streamlit Secrets
    client = genai.Client(api_key=st.secrets["GEMINI_API_KEY"])
    MY_EMAIL = st.secrets["MY_GMAIL"]
    APP_PASS = st.secrets["GMAIL_APP_PASSWORD"]
except Exception as e:
    st.error("Credential Error: Please check your Streamlit Secrets.")

# --- SCRAPER WITH SCROLLING & DATA CAPTURE ---
def run_real_hunter(niche, location, limit):
    options = Options()
    options.add_argument("--headless")
    options.add_argument("--no-sandbox")
    options.add_argument("--disable-dev-shm-usage")
    driver = webdriver.Chrome(options=options)
    
    stealth(driver, languages=["en-US", "en"], vendor="Google Inc.", platform="Win32", webgl_vendor="Intel Inc.", renderer="Intel Iris OpenGL Engine", fix_hairline=True)
    
    leads = []
    try:
        query = f"{niche} in {location}".replace(" ", "+")
        driver.get(f"https://www.google.com/maps/search/{query}")
        time.sleep(5)
        
        # SCROLL TO LOAD MORE LEADS
        try:
            # Find the scrollable side pane
            pane = driver.find_element(By.XPATH, '//div[@role="feed"]')
            for _ in range(4): # Scroll multiple times to reach the requested limit
                pane.send_keys(Keys.PAGE_DOWN)
                time.sleep(2)
        except:
            pass

        # Capture results
        elements = driver.find_elements(By.CLASS_NAME, "hfpxzc")
        for e in elements[:limit]:
            name = e.get_attribute("aria-label")
            
            # Find Rating within the same container
            try:
                parent = e.find_element(By.XPATH, "./../../..")
                rating = parent.find_element(By.CLASS_NAME, "MW4Y7c").text
            except:
                rating = "No Rating"
                
            leads.append({
                "Business Name": name, 
                "Location": location, 
                "Rating": rating, 
                "Status": "Lead Found"
            })
    except Exception as ex:
        st.error(f"Scraper Error: {ex}")
    finally:
        driver.quit()
    return leads

# --- UI NAVIGATION ---
with st.sidebar:
    st.title("Main Menu")
    page = st.radio("Navigation", ["Dashboard", "History"])

if page == "Dashboard":
    st.title("Outreach Dashboard")
    col1, col2 = st.columns(2)
    with col1:
        n_in = st.text_input("Target Niche", placeholder="e.g. Lawyers")
        l_in = st.text_input("Location", placeholder="e.g. Guwahati")
    with col2:
        num_leads = st.slider("Lead Quantity", 5, 50, 10)
        p_in = st.text_area("Pitch Focus", placeholder="Enter offer details...")

    if st.button("Initialize Hunter Machine"):
        if n_in and l_in:
            with st.spinner(f"Scraping {num_leads} leads for {n_in}..."):
                results = run_real_hunter(n_in, l_in, num_leads)
                st.session_state['leads'] = results
                st.session_state['history'].append(f"Found {len(results)} {n_in} in {l_in}")
                st.success(f"Success! Found {len(results)} leads.")
        else:
            st.warning("Please fill in both Niche and Location.")

    # --- RESULTS TABLE & REPORT ---
    if st.session_state['leads']:
        st.subheader("Lead Report")
        df = pd.DataFrame(st.session_state['leads'])
        st.table(df) # Clean table format as requested
        
        # Word Document Export
        doc = Document()
        doc.add_heading('Leads Summary Report', 0)
        table = doc.add_table(rows=1, cols=4)
        for i, h in enumerate(["Business Name", "Location", "Rating", "Status"]):
            table.rows[0].cells[i].text = h
        for lead in st.session_state['leads']:
            row = table.add_row().cells
            row[0].text, row[1].text, row[2].text, row[3].text = lead["Business Name"], lead["Location"], lead["Rating"], lead["Status"]
        
        bio = io.BytesIO()
        doc.save(bio)
        st.download_button("Download Report (Word)", data=bio.getvalue(), file_name="leads.docx")

        # PITCHING SECTION
        st.markdown("### Send Pitches")
        for idx, lead in enumerate(st.session_state['leads']):
            if st.button(f"Generate & Send Email to {lead['Business Name']}", key=f"pitch_{idx}"):
                try:
                    # THE 404 FIX: No 'models/' prefix, direct string only
                    response = client.models.generate_content(
                        model="gemini-1.5-flash", 
                        contents=f"Write a short, professional pitch for {lead['Business Name']}. Context: {p_in}"
                    )
                    
                    # Email Logic
                    msg = MIMEMultipart()
                    msg['From'], msg['To'], msg['Subject'] = MY_EMAIL, "client@example.com", "Business Proposal"
                    msg.attach(MIMEText(response.text, 'plain'))
                    
                    with smtplib.SMTP('smtp.gmail.com', 587) as server:
                        server.starttls()
                        server.login(MY_EMAIL, APP_PASS)
                        server.send_message(msg)
                    st.toast(f"Email sent to {lead['Business Name']}!")
                except Exception as e:
                    st.error(f"API Error: {e}")

elif page == "History":
    st.title("Activity History")
    for log in reversed(st.session_state['history']):
        st.write(f"• {log}")
